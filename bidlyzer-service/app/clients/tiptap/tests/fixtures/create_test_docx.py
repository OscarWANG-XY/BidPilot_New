from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_docx(output_path):
    """创建一个包含各种格式的测试文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('测试文档', 0)
    
    # 添加一级标题
    doc.add_heading('一级标题', level=1)
    doc.add_paragraph('这是一级标题下的段落')
    
    # 添加二级标题
    doc.add_heading('二级标题', level=2)
    doc.add_paragraph('这是二级标题下的段落')
    
    # 添加带格式的段落
    p = doc.add_paragraph()
    p.add_run('这是').bold = True
    p.add_run('一个包含').italic = True
    p.add_run('不同格式的段落')
    
    # 添加列表
    doc.add_heading('无序列表', level=2)
    doc.add_paragraph('列表项 1', style='List Bullet')
    doc.add_paragraph('列表项 2', style='List Bullet')
    doc.add_paragraph('列表项 3', style='List Bullet')
    
    doc.add_heading('有序列表', level=2)
    doc.add_paragraph('第一项', style='List Number')
    doc.add_paragraph('第二项', style='List Number')
    doc.add_paragraph('第三项', style='List Number')
    
    # 添加表格
    doc.add_heading('表格示例', level=2)
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # 填充表格
    cell = table.cell(0, 0)
    cell.text = "单元格 1"
    cell = table.cell(0, 1)
    cell.text = "单元格 2"
    cell = table.cell(1, 0)
    cell.text = "单元格 3"
    cell = table.cell(1, 1)
    cell.text = "单元格 4"
    
    # 保存文档
    doc.save(output_path)

if __name__ == '__main__':
    import os
    output_path = os.path.join(os.path.dirname(__file__), 'sample.docx')
    create_test_docx(output_path)
    print(f"测试文档已创建: {output_path}") 